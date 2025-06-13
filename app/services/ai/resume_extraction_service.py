import asyncio
import logging
import json
import spacy
from openai import OpenAI
from typing import Dict, Any, List, Optional
from datetime import datetime
import PyPDF2
import docx
import fitz  # PyMuPDF for better PDF extraction
from io import BytesIO
import re
from dataclasses import dataclass
from ... import db
from ...models.resume import ExtractionJob, ResumeSection, Resume
from ...config import Config

# Configure logging
logger = logging.getLogger(__name__)

@dataclass
class ExtractionProgress:
    """Data class for tracking extraction progress"""
    stage: str
    progress: int
    current_operation: str
    stage_changed: bool = False
    extracted_data: Optional[Dict] = None

class ResumeExtractionService:
    """
    Advanced AI-powered resume extraction service using NLP, NER, and OpenAI
    """
    
    def __init__(self):
        # Initialize OpenAI
        self.openai_client = OpenAI(api_key=Config.OPENAI_API_KEY)
        
        # Initialize spaCy NLP model
        try:
            self.nlp = spacy.load("en_core_web_lg")
        except OSError:
            logger.warning("spaCy model 'en_core_web_lg' not found, using 'en_core_web_sm'")
            self.nlp = spacy.load("en_core_web_sm")
        
        # Extraction progress tracking
        self.extraction_progress = {}
        
        # AI prompt templates
        self.prompt_templates = {
            'comprehensive_profile': """
            Extract comprehensive professional information from this resume text.
            
            Resume Text:
            {text}
            
            Please extract and structure the following information in JSON format:
            
            1. Personal Information:
               - Full name
               - Professional title/headline
               - Location (city, state, country)
               - Date of birth (if mentioned)
               - Nationality (if mentioned)
            
            2. Contact Information:
               - Email address(es)
               - Phone number(s)
               - LinkedIn profile
               - Website/Portfolio URL
               - GitHub profile
               - Other social media
            
            3. Professional Summary:
               - Summary/objective statement
               - Years of experience
               - Key expertise areas
               - Career highlights
            
            Return ONLY valid JSON with the extracted information. Use null for missing data.
            """,
            
            'professional_details': """
            Extract detailed professional experience and education from this resume.
            
            Resume Text:
            {text}
            
            Extract in JSON format:
            
            1. Work Experience (array of objects):
               - Job title
               - Company name
               - Company industry
               - Employment type (full-time, part-time, contract, etc.)
               - Start date (month/year)
               - End date (month/year or "Present")
               - Location
               - Job description/responsibilities
               - Key achievements
               - Technologies used
            
            2. Education (array of objects):
               - Degree type
               - Field of study
               - Institution name
               - Location
               - Graduation date
               - GPA (if mentioned)
               - Relevant coursework
               - Academic achievements
            
            3. Certifications (array of objects):
               - Certification name
               - Issuing organization
               - Issue date
               - Expiration date
               - Credential ID
               - Verification URL
            
            Return ONLY valid JSON.
            """,
            
            'skills_analysis': """
            Perform comprehensive skills analysis on this resume.
            
            Resume Text:
            {text}
            
            Extract and categorize skills in JSON format:
            
            1. Technical Skills:
               - Programming languages
               - Frameworks and libraries
               - Databases
               - Tools and software
               - Cloud platforms
               - Operating systems
            
            2. Professional Skills:
               - Industry-specific skills
               - Methodologies (Agile, Scrum, etc.)
               - Business skills
               - Leadership skills
            
            3. Soft Skills:
               - Communication
               - Problem-solving
               - Teamwork
               - Leadership
               - Other interpersonal skills
            
            4. Languages:
               - Language name
               - Proficiency level
               - Certifications
            
            For each skill category, include proficiency level if mentioned (Beginner, Intermediate, Advanced, Expert).
            
            Return ONLY valid JSON.
            """,
            
            'experience_mapping': """
            Create a detailed experience timeline and achievement mapping.
            
            Resume Text:
            {text}
            
            Extract in JSON format:
            
            1. Career Timeline:
               - Chronological work history
               - Career progression patterns
               - Industry transitions
               - Promotion history
            
            2. Project Portfolio:
               - Project name
               - Description
               - Role in project
               - Technologies used
               - Duration
               - Team size
               - Key outcomes
               - URL (if available)
            
            3. Awards and Achievements:
               - Award/recognition name
               - Awarding organization
               - Date received
               - Description
               - Impact/significance
            
            4. Publications and Speaking:
               - Publication title
               - Publication venue
               - Date published
               - Co-authors
               - Speaking engagements
               - Conference presentations
            
            Return ONLY valid JSON.
            """
        }
        
        # NER entity patterns for resume data
        self.ner_patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            'linkedin': r'linkedin\.com/in/[A-Za-z0-9-]+',
            'github': r'github\.com/[A-Za-z0-9-]+',
            'website': r'https?://(?:[-\w.])+(?:\.[a-zA-Z]{2,})+(?:/(?:[\w/_.])*)?(?:\?(?:[\w&=%.])*)?(?:#(?:\w)*)?',
            'date': r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b|\b\d{1,2}/\d{1,2}/\d{4}\b|\b\d{4}\b'
        }

    async def process_resume_extraction(self, extraction_id: str, file_path: str, options: Dict[str, Any]):
        """
        Main extraction processing method
        """
        try:
            # Initialize progress tracking
            self.extraction_progress[extraction_id] = ExtractionProgress(
                stage="Text Extraction",
                progress=0,
                current_operation="Extracting text from document..."
            )
            
            # Stage 1: Extract text from document
            await self._update_progress(extraction_id, "Text Extraction", 0, "Extracting text from document...")
            text_content = await self._extract_text_from_file(file_path)
            await self._update_progress(extraction_id, "Text Extraction", 20, "Text extraction completed")
            
            # Stage 2: NLP Processing
            await self._update_progress(extraction_id, "NLP Processing", 20, "Processing with Natural Language Processing...")
            nlp_analysis = await self._perform_nlp_analysis(text_content, options) if options.get('use_nlp') else {}
            await self._update_progress(extraction_id, "NLP Processing", 45, "NLP analysis completed")
            
            # Stage 3: NER Analysis
            await self._update_progress(extraction_id, "NER Analysis", 45, "Identifying entities with Named Entity Recognition...")
            ner_entities = await self._perform_ner_extraction(text_content, options) if options.get('use_ner') else {}
            await self._update_progress(extraction_id, "NER Analysis", 75, "NER analysis completed")
            
            # Stage 4: AI Enhancement
            await self._update_progress(extraction_id, "AI Enhancement", 75, "Enhancing data with AI prompt templates...")
            ai_extracted_data = await self._perform_ai_extraction(text_content, options) if options.get('use_ai_enhancement') else {}
            await self._update_progress(extraction_id, "AI Enhancement", 95, "AI enhancement completed")
            
            # Stage 5: Finalization
            await self._update_progress(extraction_id, "Finalization", 95, "Finalizing and structuring data...")
            final_data = await self._finalize_extraction(nlp_analysis, ner_entities, ai_extracted_data, options)
            
            # Save results to database
            await self._save_extraction_results(extraction_id, final_data)
            
            await self._update_progress(extraction_id, "Finalization", 100, "Extraction completed successfully", final_data)
            
            # Mark extraction as completed
            extraction_job = ExtractionJob.query.get(extraction_id)
            if extraction_job:
                extraction_job.status = 'completed'
                extraction_job.completed_at = datetime.utcnow()
                extraction_job.extracted_data = final_data
                
                # Update resume status
                resume = Resume.query.get(extraction_job.resume_id)
                if resume:
                    resume.status = 'extracted'
                    resume.extracted_data = final_data
                    resume.last_updated = datetime.utcnow()
                
                db.session.commit()
            
        except Exception as e:
            logger.error(f"Extraction failed for {extraction_id}: {str(e)}")
            await self._handle_extraction_error(extraction_id, str(e))

    async def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from PDF, DOC, or DOCX files"""
        try:
            file_extension = file_path.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return await self._extract_from_pdf(file_path)
            elif file_extension in ['doc', 'docx']:
                return await self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            raise

    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF for better accuracy"""
        try:
            text_content = ""
            pdf_document = fitz.open(file_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                text_content += page.get_text()
                text_content += "\n\n"  # Add page breaks
            
            pdf_document.close()
            return text_content.strip()
            
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed, falling back to PyPDF2: {str(e)}")
            # Fallback to PyPDF2
            return await self._extract_from_pdf_pypdf2(file_path)

    async def _extract_from_pdf_pypdf2(self, file_path: str) -> str:
        """Fallback PDF extraction using PyPDF2"""
        text_content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content += page.extract_text()
                text_content += "\n\n"
        return text_content.strip()

    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise

    async def _perform_nlp_analysis(self, text: str, options: Dict) -> Dict[str, Any]:
        """Perform NLP analysis using spaCy"""
        try:
            doc = self.nlp(text)
            
            analysis = {
                'sentences': [sent.text for sent in doc.sents],
                'tokens': len(doc),
                'entities': [
                    {
                        'text': ent.text,
                        'label': ent.label_,
                        'start': ent.start_char,
                        'end': ent.end_char,
                        'description': spacy.explain(ent.label_)
                    }
                    for ent in doc.ents
                ],
                'noun_phrases': [chunk.text for chunk in doc.noun_chunks],
                'pos_tags': [
                    {
                        'text': token.text,
                        'pos': token.pos_,
                        'tag': token.tag_,
                        'lemma': token.lemma_
                    }
                    for token in doc if not token.is_space
                ]
            }
            
            return analysis
            
        except Exception as e:
            logger.error(f"NLP analysis failed: {str(e)}")
            return {}

    async def _perform_ner_extraction(self, text: str, options: Dict) -> Dict[str, Any]:
        """Perform Named Entity Recognition using regex patterns and spaCy"""
        try:
            entities = {}
            
            # Extract using regex patterns
            for entity_type, pattern in self.ner_patterns.items():
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    entities[entity_type] = matches
            
            # Extract using spaCy NER
            doc = self.nlp(text)
            spacy_entities = {}
            
            for ent in doc.ents:
                entity_label = ent.label_.lower()
                if entity_label not in spacy_entities:
                    spacy_entities[entity_label] = []
                spacy_entities[entity_label].append({
                    'text': ent.text,
                    'confidence': 1.0,  # spaCy doesn't provide confidence scores by default
                    'start': ent.start_char,
                    'end': ent.end_char
                })
            
            return {
                'regex_entities': entities,
                'nlp_entities': spacy_entities
            }
            
        except Exception as e:
            logger.error(f"NER extraction failed: {str(e)}")
            return {}

    async def _perform_ai_extraction(self, text: str, options: Dict) -> Dict[str, Any]:
        """Perform AI-powered extraction using OpenAI GPT models"""
        try:
            extracted_data = {}
            templates = options.get('extraction_templates', [])
            
            for template_name in templates:
                if template_name in self.prompt_templates:
                    logger.info(f"Running AI extraction with template: {template_name}")
                    
                    prompt = self.prompt_templates[template_name].format(text=text)
                    
                    response = await self._call_openai_api(prompt)
                    
                    if response:
                        try:
                            # Parse JSON response
                            template_data = json.loads(response)
                            extracted_data[template_name] = template_data
                        except json.JSONDecodeError as e:
                            logger.error(f"Failed to parse JSON from template {template_name}: {str(e)}")
                            extracted_data[template_name] = {'raw_response': response}
            
            # Process custom prompts if provided
            custom_prompts = options.get('custom_prompts', {})
            for prompt_name, prompt_text in custom_prompts.items():
                response = await self._call_openai_api(prompt_text.format(text=text))
                if response:
                    extracted_data[f'custom_{prompt_name}'] = response
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"AI extraction failed: {str(e)}")
            return {}

    async def _call_openai_api(self, prompt: str) -> Optional[str]:
        """Make API call to OpenAI"""
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system", 
                        "content": "You are an expert resume parser. Extract information accurately and return only valid JSON when requested."
                    },
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.1,
                timeout=30
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.error(f"OpenAI API call failed: {str(e)}")
            return None

    async def _finalize_extraction(self, nlp_data: Dict, ner_data: Dict, ai_data: Dict, options: Dict) -> Dict[str, Any]:
        """Combine and finalize all extracted data"""
        try:
            final_data = {
                'personal_information': {},
                'contact_details': {},
                'professional_summary': {},
                'work_experience': [],
                'education': [],
                'skills': {},
                'certifications': [],
                'languages': [],
                'projects': [],
                'awards_and_achievements': [],
                'metadata': {
                    'extraction_timestamp': datetime.utcnow().isoformat(),
                    'extraction_method': 'ai_nlp_ner_combined',
                    'confidence_scores': {},
                    'data_sources': {
                        'nlp_analysis': bool(nlp_data),
                        'ner_extraction': bool(ner_data),
                        'ai_enhancement': bool(ai_data)
                    }
                }
            }
            
            # Merge AI-extracted data
            for template_name, template_data in ai_data.items():
                if template_name == 'comprehensive_profile' and isinstance(template_data, dict):
                    if 'Personal Information' in template_data:
                        final_data['personal_information'].update(template_data['Personal Information'])
                    if 'Contact Information' in template_data:
                        final_data['contact_details'].update(template_data['Contact Information'])
                    if 'Professional Summary' in template_data:
                        final_data['professional_summary'].update(template_data['Professional Summary'])
                
                elif template_name == 'professional_details' and isinstance(template_data, dict):
                    if 'Work Experience' in template_data:
                        final_data['work_experience'] = template_data['Work Experience']
                    if 'Education' in template_data:
                        final_data['education'] = template_data['Education']
                    if 'Certifications' in template_data:
                        final_data['certifications'] = template_data['Certifications']
                
                elif template_name == 'skills_analysis' and isinstance(template_data, dict):
                    final_data['skills'].update(template_data)
                
                elif template_name == 'experience_mapping' and isinstance(template_data, dict):
                    if 'Project Portfolio' in template_data:
                        final_data['projects'] = template_data['Project Portfolio']
                    if 'Awards and Achievements' in template_data:
                        final_data['awards_and_achievements'] = template_data['Awards and Achievements']
            
            # Enhance with NER data
            if ner_data.get('regex_entities'):
                regex_entities = ner_data['regex_entities']
                
                if 'email' in regex_entities:
                    final_data['contact_details']['email'] = regex_entities['email'][0] if regex_entities['email'] else None
                
                if 'phone' in regex_entities:
                    final_data['contact_details']['phone'] = regex_entities['phone'][0] if regex_entities['phone'] else None
                
                if 'linkedin' in regex_entities:
                    final_data['contact_details']['linkedin'] = f"https://{regex_entities['linkedin'][0]}" if regex_entities['linkedin'] else None
                
                if 'github' in regex_entities:
                    final_data['contact_details']['github'] = f"https://{regex_entities['github'][0]}" if regex_entities['github'] else None
            
            # Calculate confidence scores
            final_data['metadata']['confidence_scores'] = self._calculate_confidence_scores(final_data, nlp_data, ner_data, ai_data)
            
            return final_data
            
        except Exception as e:
            logger.error(f"Data finalization failed: {str(e)}")
            return {'error': f'Finalization failed: {str(e)}'}

    def _calculate_confidence_scores(self, final_data: Dict, nlp_data: Dict, ner_data: Dict, ai_data: Dict) -> Dict[str, float]:
        """Calculate confidence scores for each section"""
        confidence_scores = {}
        
        # Calculate based on data completeness and source reliability
        for section_name, section_data in final_data.items():
            if section_name == 'metadata':
                continue
                
            score = 0.0
            
            # Base score from data presence
            if section_data:
                score += 0.4
                
                # Bonus for AI extraction
                if ai_data:
                    score += 0.3
                
                # Bonus for NER confirmation
                if ner_data:
                    score += 0.2
                
                # Bonus for NLP analysis
                if nlp_data:
                    score += 0.1
            
            confidence_scores[section_name] = min(score, 1.0)
        
        return confidence_scores

    async def _save_extraction_results(self, extraction_id: str, data: Dict[str, Any]):
        """Save extraction results to database"""
        try:
            extraction_job = ExtractionJob.query.get(extraction_id)
            if not extraction_job:
                return
            
            resume_id = extraction_job.resume_id
            
            # Save each section
            for section_name, section_data in data.items():
                if section_name == 'metadata':
                    continue
                
                confidence_score = data.get('metadata', {}).get('confidence_scores', {}).get(section_name, 0.0)
                
                # Check if section already exists
                existing_section = ResumeSection.query.filter_by(
                    resume_id=resume_id,
                    section_name=section_name
                ).first()
                
                if existing_section:
                    existing_section.section_data = section_data
                    existing_section.confidence_score = confidence_score
                    existing_section.last_updated = datetime.utcnow()
                else:
                    section = ResumeSection(
                        resume_id=resume_id,
                        section_name=section_name,
                        section_data=section_data,
                        confidence_score=confidence_score,
                        extraction_method='ai_nlp_ner',
                        last_updated=datetime.utcnow()
                    )
                    db.session.add(section)
            
            db.session.commit()
            
        except Exception as e:
            logger.error(f"Failed to save extraction results: {str(e)}")

    async def _update_progress(self, extraction_id: str, stage: str, progress: int, operation: str, data: Optional[Dict] = None):
        """Update extraction progress"""
        current_progress = self.extraction_progress.get(extraction_id)
        stage_changed = current_progress is None or current_progress.stage != stage
        
        self.extraction_progress[extraction_id] = ExtractionProgress(
            stage=stage,
            progress=progress,
            current_operation=operation,
            stage_changed=stage_changed,
            extracted_data=data
        )
        
        # Update database
        extraction_job = ExtractionJob.query.get(extraction_id)
        if extraction_job:
            extraction_job.progress = progress
            extraction_job.current_stage = stage
            extraction_job.status = 'processing' if progress < 100 else 'completed'
            db.session.commit()

    async def _handle_extraction_error(self, extraction_id: str, error_message: str):
        """Handle extraction errors"""
        self.extraction_progress[extraction_id] = ExtractionProgress(
            stage="Error",
            progress=0,
            current_operation=f"Extraction failed: {error_message}"
        )
        
        # Update database
        extraction_job = ExtractionJob.query.get(extraction_id)
        if extraction_job:
            extraction_job.status = 'failed'
            extraction_job.error_message = error_message
            extraction_job.completed_at = datetime.utcnow()
            
            # Update resume status
            resume = Resume.query.get(extraction_job.resume_id)
            if resume:
                resume.status = 'extraction_failed'
            
            db.session.commit()

    def get_extraction_progress(self, extraction_id: str) -> Dict[str, Any]:
        """Get current extraction progress"""
        progress = self.extraction_progress.get(extraction_id)
        
        if not progress:
            # Try to get from database
            extraction_job = ExtractionJob.query.get(extraction_id)
            if extraction_job:
                return {
                    'current_stage': extraction_job.current_stage or 'Unknown',
                    'overall_progress': extraction_job.progress or 0,
                    'stage_progress': extraction_job.progress or 0,
                    'current_operation': f"Status: {extraction_job.status}",
                    'stage_changed': False
                }
            return {
                'current_stage': 'Unknown',
                'overall_progress': 0,
                'stage_progress': 0,
                'current_operation': 'Progress information not available',
                'stage_changed': False
            }
        
        return {
            'current_stage': progress.stage,
            'overall_progress': progress.progress,
            'stage_progress': progress.progress,
            'current_operation': progress.current_operation,
            'stage_changed': progress.stage_changed
        } 