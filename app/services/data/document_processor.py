import os
import hashlib
import logging
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
import PyPDF2
import pdfplumber
from docx import Document
import spacy
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.tag import pos_tag
from nltk.chunk import ne_chunk
from bs4 import BeautifulSoup
import re

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger')

try:
    nltk.data.find('chunkers/maxent_ne_chunker')
except LookupError:
    nltk.download('maxent_ne_chunker')

try:
    nltk.data.find('corpora/words')
except LookupError:
    nltk.download('words')

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Comprehensive document processing service with AI capabilities"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.html']
        self.nlp_model = None
        self.embedding_model = None
        self.stop_words = set(stopwords.words('english'))
        
        # Initialize models
        self._initialize_models()
    
    def _initialize_models(self):
        """Initialize NLP and embedding models"""
        try:
            # Load spaCy model (download with: python -m spacy download en_core_web_sm)
            self.nlp_model = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy English model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp_model = None
        
        try:
            # Load sentence transformer model
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        except Exception as e:
            logger.warning(f"Failed to load embedding model: {e}")
            self.embedding_model = None
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from various file formats
        
        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, txt, html)
            
        Returns:
            Tuple of (extracted_text, extraction_metadata)
        """
        metadata = {
            'extraction_method': None,
            'pages': 0,
            'word_count': 0,
            'character_count': 0,
            'extraction_time': None,
            'errors': []
        }
        
        start_time = datetime.now()
        
        try:
            if file_type.lower() == 'pdf':
                text = self._extract_from_pdf(file_path, metadata)
            elif file_type.lower() == 'docx':
                text = self._extract_from_docx(file_path, metadata)
            elif file_type.lower() == 'txt':
                text = self._extract_from_txt(file_path, metadata)
            elif file_type.lower() == 'html':
                text = self._extract_from_html(file_path, metadata)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Calculate metadata
            metadata['word_count'] = len(text.split())
            metadata['character_count'] = len(text)
            metadata['extraction_time'] = (datetime.now() - start_time).total_seconds()
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            metadata['errors'].append(str(e))
            return "", metadata
    
    def _extract_from_pdf(self, file_path: str, metadata: Dict[str, Any]) -> str:
        """Extract text from PDF using both PyPDF2 and pdfplumber"""
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                metadata['extraction_method'] = 'pdfplumber'
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    metadata['extraction_method'] = 'PyPDF2'
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n"
                        
            except Exception as e2:
                metadata['errors'].append(f"Both PDF extraction methods failed: pdfplumber: {e}, PyPDF2: {e2}")
                raise e2
        
        return text.strip()
    
    def _extract_from_docx(self, file_path: str, metadata: Dict[str, Any]) -> str:
        """Extract text from DOCX files"""
        metadata['extraction_method'] = 'python-docx'
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    def _extract_from_txt(self, file_path: str, metadata: Dict[str, Any]) -> str:
        """Extract text from plain text files"""
        metadata['extraction_method'] = 'plain_text'
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    
    def _extract_from_html(self, file_path: str, metadata: Dict[str, Any]) -> str:
        """Extract text from HTML files"""
        metadata['extraction_method'] = 'beautifulsoup'
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            html_content = file.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def analyze_content(self, text: str) -> Dict[str, Any]:
        """
        Perform comprehensive NLP analysis on the extracted text
        
        Args:
            text: The text to analyze
            
        Returns:
            Dictionary containing analysis results
        """
        analysis = {
            'summary_stats': self._get_summary_stats(text),
            'entities': self._extract_entities(text),
            'keywords': self._extract_keywords(text),
            'sentiment': self._analyze_sentiment(text),
            'topics': self._extract_topics(text),
            'readability': self._calculate_readability(text),
            'content_structure': self._analyze_structure(text)
        }
        
        return analysis
    
    def _get_summary_stats(self, text: str) -> Dict[str, Any]:
        """Get basic statistics about the text"""
        sentences = sent_tokenize(text)
        words = word_tokenize(text)
        
        return {
            'total_characters': len(text),
            'total_words': len(words),
            'total_sentences': len(sentences),
            'avg_words_per_sentence': len(words) / len(sentences) if sentences else 0,
            'avg_characters_per_word': len(text) / len(words) if words else 0
        }
    
    def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities using both NLTK and spaCy"""
        entities = []
        
        # NLTK NER
        try:
            sentences = sent_tokenize(text)
            for sentence in sentences[:5]:  # Limit to first 5 sentences for performance
                tokens = word_tokenize(sentence)
                pos_tags = pos_tag(tokens)
                chunks = ne_chunk(pos_tags, binary=False)
                
                for chunk in chunks:
                    if hasattr(chunk, 'label'):
                        entity_text = ' '.join([token for token, pos in chunk.leaves()])
                        entities.append({
                            'text': entity_text,
                            'label': chunk.label(),
                            'method': 'nltk'
                        })
        except Exception as e:
            logger.warning(f"NLTK NER failed: {e}")
        
        # spaCy NER (if available)
        if self.nlp_model:
            try:
                # Process in chunks to avoid memory issues
                max_length = 1000000  # spaCy's default max length
                text_chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
                
                for chunk in text_chunks:
                    doc = self.nlp_model(chunk)
                    for ent in doc.ents:
                        entities.append({
                            'text': ent.text,
                            'label': ent.label_,
                            'confidence': float(ent._.get('confidence', 0.0)) if hasattr(ent._, 'confidence') else None,
                            'method': 'spacy'
                        })
            except Exception as e:
                logger.warning(f"spaCy NER failed: {e}")
        
        # Remove duplicates and sort by frequency
        unique_entities = {}
        for entity in entities:
            key = (entity['text'].lower(), entity['label'])
            if key not in unique_entities:
                unique_entities[key] = entity
                unique_entities[key]['frequency'] = 1
            else:
                unique_entities[key]['frequency'] += 1
        
        return sorted(unique_entities.values(), key=lambda x: x['frequency'], reverse=True)
    
    def _extract_keywords(self, text: str) -> List[Dict[str, Any]]:
        """Extract keywords using TF-IDF and frequency analysis"""
        words = word_tokenize(text.lower())
        
        # Filter out stopwords and non-alphabetic tokens
        filtered_words = [
            word for word in words 
            if word.isalpha() and word not in self.stop_words and len(word) > 2
        ]
        
        # Calculate word frequencies
        word_freq = {}
        for word in filtered_words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top keywords by frequency
        total_words = len(filtered_words)
        keywords = []
        
        for word, freq in sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]:
            keywords.append({
                'word': word,
                'frequency': freq,
                'tf_score': freq / total_words
            })
        
        return keywords
    
    def _analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """Basic sentiment analysis using keyword matching"""
        positive_words = ['good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'positive', 'success', 'achieve', 'accomplish']
        negative_words = ['bad', 'terrible', 'awful', 'horrible', 'negative', 'fail', 'failure', 'problem', 'issue', 'difficult']
        
        words = word_tokenize(text.lower())
        
        positive_count = sum(1 for word in words if word in positive_words)
        negative_count = sum(1 for word in words if word in negative_words)
        
        total_sentiment_words = positive_count + negative_count
        
        if total_sentiment_words == 0:
            sentiment_score = 0.0
            sentiment_label = 'neutral'
        else:
            sentiment_score = (positive_count - negative_count) / total_sentiment_words
            if sentiment_score > 0.1:
                sentiment_label = 'positive'
            elif sentiment_score < -0.1:
                sentiment_label = 'negative'
            else:
                sentiment_label = 'neutral'
        
        return {
            'score': sentiment_score,
            'label': sentiment_label,
            'positive_words': positive_count,
            'negative_words': negative_count
        }
    
    def _extract_topics(self, text: str) -> List[str]:
        """Extract potential topics using keyword clustering"""
        # This is a simplified topic extraction
        # In a production system, you might use LDA or other topic modeling techniques
        
        keywords = self._extract_keywords(text)
        
        # Group related keywords (simplified approach)
        topics = []
        high_freq_keywords = [kw['word'] for kw in keywords[:10] if kw['frequency'] > 2]
        
        if high_freq_keywords:
            topics.append(' '.join(high_freq_keywords[:5]))
        
        return topics
    
    def _calculate_readability(self, text: str) -> Dict[str, Any]:
        """Calculate basic readability metrics"""
        sentences = sent_tokenize(text)
        words = word_tokenize(text)
        
        # Count syllables (simplified)
        def count_syllables(word):
            vowels = 'aeiouy'
            syllables = 0
            prev_was_vowel = False
            
            for char in word.lower():
                if char in vowels:
                    if not prev_was_vowel:
                        syllables += 1
                    prev_was_vowel = True
                else:
                    prev_was_vowel = False
            
            return max(1, syllables)
        
        total_syllables = sum(count_syllables(word) for word in words if word.isalpha())
        
        # Simple readability metrics
        avg_sentence_length = len(words) / len(sentences) if sentences else 0
        avg_syllables_per_word = total_syllables / len(words) if words else 0
        
        return {
            'avg_sentence_length': avg_sentence_length,
            'avg_syllables_per_word': avg_syllables_per_word,
            'total_syllables': total_syllables
        }
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze the structure of the document"""
        lines = text.split('\n')
        
        # Count different types of content
        bullet_points = len([line for line in lines if re.match(r'^\s*[-*•]\s', line)])
        numbered_items = len([line for line in lines if re.match(r'^\s*\d+\.\s', line)])
        
        # Find potential headers (lines that are short and followed by longer content)
        potential_headers = []
        for i, line in enumerate(lines):
            if (len(line.strip()) < 50 and 
                len(line.strip()) > 0 and 
                i < len(lines) - 1 and 
                len(lines[i + 1].strip()) > len(line.strip())):
                potential_headers.append(line.strip())
        
        return {
            'total_lines': len(lines),
            'empty_lines': len([line for line in lines if not line.strip()]),
            'bullet_points': bullet_points,
            'numbered_items': numbered_items,
            'potential_headers': potential_headers[:10]  # Limit to first 10
        }
    
    def generate_embeddings(self, text: str) -> Optional[List[float]]:
        """Generate embeddings for the text using sentence transformers"""
        if not self.embedding_model:
            logger.warning("Embedding model not available")
            return None
        
        try:
            # Truncate text if too long
            max_length = 512  # Typical transformer limit
            if len(text) > max_length:
                text = text[:max_length]
            
            embeddings = self.embedding_model.encode(text)
            return embeddings.tolist()
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            return None
    
    def calculate_content_hash(self, text: str) -> str:
        """Calculate a hash of the content for change detection"""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()
    
    def is_supported_format(self, file_extension: str) -> bool:
        """Check if the file format is supported"""
        return file_extension.lower() in self.supported_formats 