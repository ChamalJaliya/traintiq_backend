import logging
from typing import List, Tuple, Optional
import io
import base64
from pathlib import Path

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessingService:
    """
    Service for processing uploaded documents (PDF, DOCX) and extracting text content.
    """
    
    def __init__(self):
        self.supported_formats = []
        if PDF_AVAILABLE:
            self.supported_formats.extend(['.pdf'])
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['.docx', '.doc'])
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if the file format is supported."""
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.supported_formats
    
    def extract_text_from_base64(self, base64_content: str, filename: str) -> str:
        """
        Extract text from base64 encoded file content.
        
        Args:
            base64_content: Base64 encoded file content
            filename: Original filename to determine format
            
        Returns:
            Extracted text content
        """
        try:
            # Decode base64 content
            file_content = base64.b64decode(base64_content)
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise ValueError(f"Failed to extract text from {filename}: {e}")
    
    def extract_text_from_files(self, files_data: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
        """
        Extract text from multiple files.
        
        Args:
            files_data: List of tuples (filename, base64_content)
            
        Returns:
            List of tuples (filename, extracted_text)
        """
        results = []
        for filename, base64_content in files_data:
            try:
                if self.is_supported_format(filename):
                    text_content = self.extract_text_from_base64(base64_content, filename)
                    results.append((filename, text_content))
                else:
                    logger.warning(f"Unsupported file format: {filename}")
                    results.append((filename, f"[Unsupported file format: {Path(filename).suffix}]"))
            except Exception as e:
                logger.error(f"Failed to process {filename}: {e}")
                results.append((filename, f"[Error processing file: {e}]"))
        
        return results
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF content."""
        if not PDF_AVAILABLE:
            raise ValueError("PDF processing libraries not available. Please install PyPDF2 and pdfplumber.")
        
        text_parts = []
        
        try:
            # Try with pdfplumber first (better text extraction)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            try:
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {e2}")
                raise ValueError(f"Failed to extract text from PDF: {e2}")
        
        if not text_parts:
            return "[No text content found in PDF]"
        
        return "\n\n".join(text_parts)
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX content."""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing library not available. Please install python-docx.")
        
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                return "[No text content found in DOCX]"
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return self.supported_formats.copy()
    
    def validate_file_size(self, base64_content: str, max_size_mb: int = 10) -> bool:
        """
        Validate file size from base64 content.
        
        Args:
            base64_content: Base64 encoded file content
            max_size_mb: Maximum file size in MB
            
        Returns:
            True if file size is acceptable
        """
        try:
            # Calculate approximate file size from base64
            # Base64 encoding increases size by ~33%
            estimated_size_bytes = len(base64_content) * 3 / 4
            max_size_bytes = max_size_mb * 1024 * 1024
            
            return estimated_size_bytes <= max_size_bytes
        except Exception:
            return False 